import io
import math
import pandas as pd
import streamlit as st

st.set_page_config(page_title="Internal Exam Seating Arrangement", layout="wide")

st.markdown("<h1 style='text-align: center; color: #003366;'>VVP Institute of Engineering and Technology, Solapur</h1>", unsafe_allow_html=True)
st.markdown("<h2 style='text-align: center;'>Internal Exam Seating Arrangement</h2>", unsafe_allow_html=True)
st.markdown("---")
st.markdown("Create bench-wise seating with different-class pairs (FY/TY vs SY/B.Tech)")

st.markdown("<p style='text-align: center; color: #555; font-weight: bold;'>Prepared by: Prof. Amir M. Usman Wagdarikar</p>", unsafe_allow_html=True)

st.sidebar.header("1. Upload and settings")
uploaded_file = st.sidebar.file_uploader("Upload student CSV", type=["csv"])

st.sidebar.markdown("Expected CSV columns: `StudentID`, `Name`, `Class` (FY, SY, TY, B.Tech), `Branch`")

st.sidebar.subheader("2. Classroom configuration")
num_classrooms = st.sidebar.number_input("Number of classrooms", min_value=1, value=4, step=1)

auto_assign = st.sidebar.checkbox("Auto-assign extra benches for remaining students", value=True)

# Classroom bench configuration
st.sidebar.write("**Specify benches for each classroom:**")
classroom_benches = {}
for i in range(int(num_classrooms)):
    classroom_benches[f"Classroom_{i+1}"] = st.sidebar.number_input(
        f"Classroom {i+1} - Number of benches",
        min_value=1,
        value=30,
        step=1,
        key=f"classroom_{i+1}_benches"
    )


if uploaded_file is None:
    st.warning("Please upload a student CSV to generate seating.")
    st.stop()

try:
    df = pd.read_csv(uploaded_file)
    df.columns = df.columns.str.strip()  # Strip whitespace from column names
except Exception as e:
    st.error(f"Failed to read CSV: {e}")
    st.stop()

required_cols = {"StudentID", "Name", "Class", "Branch"}
if not required_cols.issubset(df.columns):
    st.error(f"CSV required columns: {required_cols}. Found: {set(df.columns)}")
    st.stop()

# Normalize class labels
df["Class"] = df["Class"].astype(str).str.strip().str.upper().replace({"BTECH": "B.TECH", "B TECH": "B.TECH", "B.TECH": "B.TECH"})

# Normalize branch labels
df["Branch"] = df["Branch"].astype(str).str.strip().str.upper()

valid_class_map = {
    "FY": "FY",
    "SY": "SY",
    "TY": "TY",
    "B.TECH": "B.Tech",
    "BTECH": "B.Tech",
    "B TECH": "B.Tech",
}

# Keep only recognized classes
df["ClassGroup"] = df["Class"].map(valid_class_map)
invalid_rows = df[df["ClassGroup"].isna()]
if len(invalid_rows) > 0:
    st.warning("Some rows have unrecognized classes and will be ignored:")
    st.dataframe(invalid_rows)
    df = df[df["ClassGroup"].notna()].copy()

if df.empty:
    st.error("No valid students after class normalization.")
    st.stop()

# categorical order for output
order_map = {"FY": 0, "SY": 1, "TY": 2, "B.Tech": 3}

df["ClassOrder"] = df["ClassGroup"].map(order_map)

# Sort input as asked
sorted_df = df.sort_values(["Branch", "ClassOrder", "StudentID"]).reset_index(drop=True)

st.subheader("Student list (sorted by branch, class and student ID)")
st.dataframe(sorted_df[["StudentID", "Name", "Branch", "ClassGroup"]])

# Add branch-roll display
sorted_df["Branch-Roll"] = sorted_df["Branch"] + " " + sorted_df["ClassGroup"] + "-" + sorted_df["StudentID"].astype(str)
st.dataframe(sorted_df[["Branch-Roll", "Name"]])

# Build academic class pools (FY, SY, TY, B.Tech) with branches
prioritized_classes = ["FY", "SY", "TY", "B.Tech"]
all_classes = sorted_df["ClassGroup"].unique().tolist()
ordered_classes = [c for c in prioritized_classes if c in all_classes] + [c for c in all_classes if c not in prioritized_classes]

class_pools = {}
for cl in ordered_classes:
    class_pools[cl] = {}
    for branch in sorted_df[sorted_df["ClassGroup"] == cl]["Branch"].unique():
        class_pools[cl][branch] = sorted_df[(sorted_df["ClassGroup"] == cl) & (sorted_df["Branch"] == branch)].to_dict(orient="records")

pair_map = {
    "FY": "TY",
    "TY": "FY",
    "SY": "B.Tech",
    "B.Tech": "SY",
}

# Allocation algorithm
benches = []
bench_no = 1
classroom_no = 1

# Create list of (classroom_name, bench_count) pairs
classroom_list = [(name, classroom_benches[name]) for name in sorted(classroom_benches.keys())]

for classroom_name, bench_count in classroom_list:
    for bench_idx in range(int(bench_count)):
        # Find student for bench position 1
        s1 = None
        for cl in ordered_classes:
            for branch in sorted(class_pools[cl].keys()):
                if class_pools[cl][branch]:
                    s1 = class_pools[cl][branch].pop(0)
                    break
            if s1:
                break
        
        if s1 is None:
            continue
        
        # Find partner for position 2 (from different class if possible)
        s1_class = s1.get("ClassGroup", "")
        partner_class = pair_map.get(s1_class)
        s2 = None
        
        if partner_class and class_pools.get(partner_class):
            for branch in sorted(class_pools[partner_class].keys()):
                if class_pools[partner_class][branch]:
                    s2 = class_pools[partner_class][branch].pop(0)
                    break
        else:
            # Fallback: find from any other class
            for fallback_cl in ordered_classes:
                if fallback_cl != s1_class:
                    for branch in sorted(class_pools[fallback_cl].keys()):
                        if class_pools[fallback_cl][branch]:
                            s2 = class_pools[fallback_cl][branch].pop(0)
                            break
                    if s2:
                        break
        
        benches.append({
            "Classroom": classroom_name,
            "Bench": bench_idx + 1,
            "Student1": f"{s1.get('Branch', '')} {s1.get('ClassGroup', '')}-{s1.get('StudentID', '')}",
            "Branch1": s1.get('Branch', ''),
            "Class1": s1_class,
            "Student2": f"{s2.get('Branch', '')} {s2.get('ClassGroup', '')}-{s2.get('StudentID', '')}" if s2 else "(empty)",
            "Branch2": s2.get('Branch', '') if s2 else '',
            "Class2": s2.get("ClassGroup", "") if s2 else "",
        })
        
        bench_no += 1

# If auto_assign, create overflow classrooms for remaining students
if auto_assign:
    overflow_classroom_idx = 1
    while any(any(class_pools[c][b] for b in class_pools[c]) for c in ordered_classes):
        # Find class with most remaining students
        max_class = max(ordered_classes, key=lambda c: sum(len(class_pools[c][b]) for b in class_pools[c]))
        if sum(len(class_pools[max_class][b]) for b in class_pools[max_class]) == 0:
            break
        
        s1 = None
        for branch in sorted(class_pools[max_class].keys()):
            if class_pools[max_class][branch]:
                s1 = class_pools[max_class][branch].pop(0)
                break
        
        partner_class = pair_map.get(s1.get("ClassGroup", ""))
        s2 = None
        
        if partner_class and class_pools.get(partner_class):
            for branch in sorted(class_pools[partner_class].keys()):
                if class_pools[partner_class][branch]:
                    s2 = class_pools[partner_class][branch].pop(0)
                    break
        else:
            for fallback_cl in ordered_classes:
                if fallback_cl != s1.get("ClassGroup", ""):
                    for branch in sorted(class_pools[fallback_cl].keys()):
                        if class_pools[fallback_cl][branch]:
                            s2 = class_pools[fallback_cl][branch].pop(0)
                            break
                    if s2:
                        break
        
        benches.append({
            "Classroom": f"Overflow_{overflow_classroom_idx}",
            "Bench": len([b for b in benches if b["Classroom"] == f"Overflow_{overflow_classroom_idx}"]) + 1,
            "Student1": f"{s1.get('Branch', '')} {s1.get('ClassGroup', '')}-{s1.get('StudentID', '')}",
            "Branch1": s1.get('Branch', ''),
            "Class1": s1.get("ClassGroup", ""),
            "Student2": f"{s2.get('Branch', '')} {s2.get('ClassGroup', '')}-{s2.get('StudentID', '')}" if s2 else "(empty)",
            "Branch2": s2.get('Branch', '') if s2 else '',
            "Class2": s2.get("ClassGroup", "") if s2 else "",
        })
        
        # Increment overflow classroom if current has 2 more students
        if len([b for b in benches if b["Classroom"] == f"Overflow_{overflow_classroom_idx}"]) % 30 == 0:
            overflow_classroom_idx += 1

# Collect unassigned students
unassigned_records = []
for cl in ordered_classes:
    for branch in class_pools[cl]:
        unassigned_records.extend(class_pools[cl][branch])

seating_df = pd.DataFrame(benches)

st.subheader("Seating arrangement (by classroom)")

# Group by classroom and display each separately
classroom_groups = seating_df.groupby("Classroom")
for classroom_name, group_df in classroom_groups:
    st.write(f"### {classroom_name}")
    st.dataframe(group_df[["Bench", "Student1", "Branch1", "Class1", "Student2", "Branch2", "Class2"]])

if unassigned_records:
    unassigned_df = pd.DataFrame(unassigned_records)
    st.warning(f"{len(unassigned_df)} students remain unassigned. Enable auto-assign or add more classrooms.")
    st.dataframe(unassigned_df[["StudentID", "Name", "Branch", "ClassGroup"]])
else:
    st.success("All students assigned to classrooms and benches.")

# Download Word and PDF
try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_FPDF = True

    class CustomPDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 14)
            self.set_text_color(0, 51, 102)
            self.cell(0, 10, 'VVP Institute of Engineering and Technology, Solapur', ln=True, align='C')
            self.set_font('Arial', 'B', 12)
            self.set_text_color(0, 0, 0)
            self.cell(0, 8, 'Internal Exam Seating Arrangement', ln=True, align='C')
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.set_text_color(128)
            self.cell(0, 10, f'Page {self.page_no()}', align='C')

except ImportError:
    HAS_FPDF = False

if HAS_DOCX:
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
        
        doc_bytes = io.BytesIO()
        doc = Document()

        # Set header for all pages
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = 'VVP Institute of Engineering and Technology, Solapur'
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Document title
        title = doc.add_heading('Internal Exam Seating Arrangement', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        classroom_groups_word = seating_df.groupby("Classroom")
        for classroom_name, group_df in classroom_groups_word:
            doc.add_heading(f'Classroom: {classroom_name}', level=2)
            table = doc.add_table(rows=1, cols=7)
            hdr_cells = table.rows[0].cells
            headers = ["Bench", "Student 1", "Branch 1", "Class 1", "Student 2", "Branch 2", "Class 2"]
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            
            # Set row height for header
            table.rows[0].height = Inches(0.4)
            
            for _, row in group_df.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row.Bench)
                cells[1].text = str(row.Student1)
                cells[2].text = str(row.Branch1)
                cells[3].text = str(row.Class1)
                cells[4].text = str(row.Student2)
                cells[5].text = str(row.Branch2)
                cells[6].text = str(row.Class2)
                # Set row height for data rows
                table.rows[-1].height = Inches(0.4)
            
            doc.add_paragraph("")  # Add space between classrooms
        
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        st.download_button("⬇️ Download arrangement as Word (.docx)", data=doc_bytes, file_name="seating_arrangement.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
else:
    st.warning("⚠️ Word download unavailable. Install: `pip install python-docx`")

if FPDF is not None:
    try:
        pdf = CustomPDF(orientation='P', unit='mm', format='A4')
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        classroom_groups_pdf = seating_df.groupby("Classroom")
        for classroom_name, group_df in classroom_groups_pdf:
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 8, f'Classroom: {classroom_name}', ln=True)
            pdf.ln(2)
            
            pdf.set_font('Arial', 'B', 10)
            col_widths = [10, 30, 20, 15, 30, 20, 15]
            headers = ["Bench", "Student1", "Branch1", "Class1", "Student2", "Branch2", "Class2"]
            for i, header in enumerate(headers):
                pdf.cell(col_widths[i], 10, header, border=1)
            pdf.ln()
            
            pdf.set_font('Arial', '', 9)
            for _, row in group_df.iterrows():
                pdf.cell(col_widths[0], 10, str(row.Bench), border=1)
                pdf.cell(col_widths[1], 10, str(row.Student1), border=1)
                pdf.cell(col_widths[2], 10, str(row.Branch1), border=1)
                pdf.cell(col_widths[3], 10, str(row.Class1), border=1)
                pdf.cell(col_widths[4], 10, str(row.Student2), border=1)
                pdf.cell(col_widths[5], 10, str(row.Branch2), border=1)
                pdf.cell(col_widths[6], 10, str(row.Class2), border=1)
                pdf.ln()
            
            pdf.ln(5)
        
        pdf_bytes = pdf.output()
        if isinstance(pdf_bytes, bytearray):
            pdf_bytes = bytes(pdf_bytes)
        st.download_button("⬇️ Download arrangement as PDF", data=pdf_bytes, file_name="seating_arrangement.pdf", mime="application/pdf")
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
else:
    st.warning("Install 'fpdf2' to enable PDF download: pip install fpdf2")

st.success("Seating generation complete!")
